import os
from typing import Optional, Dict, List, Union, Any, Callable
from pathlib import Path
import PyPDF2
from docx import Document
import pandas as pd
from tqdm import tqdm
import re
import logging
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
import time
import hashlib
import docx
import pdfplumber
import pytesseract
from PIL import Image
import io
import aiofiles
import asyncio
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from .vector_store import DocumentVectorStore


class DocumentProcessor:
    """Processes documents and extracts text with support for tables and images."""
    
    def __init__(
        self,
        upload_dir: Union[str, Path] = "uploads",
        skip_images: bool = True,
        extract_tables: bool = True,
        logging_level: int = logging.INFO
    ):
        """
        Initialize the document processor.
        
        Args:
            upload_dir: Directory for uploaded files
            skip_images: Whether to skip image processing
            extract_tables: Whether to extract tables
            logging_level: Logging level
        """
        # Configure logging
        logging.basicConfig(level=logging_level)
        self.logger = logging.getLogger(__name__)
        
        # Initialize parameters
        self.upload_dir = Path(upload_dir)
        self.skip_images = skip_images
        self.extract_tables = extract_tables
        
        # Create upload directory
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def _get_file_hash(self, file_path: Union[str, Path]) -> str:
        """Generate a hash for a file."""
        try:
            with open(file_path, "rb") as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception as e:
            self.logger.error(f"Error generating file hash: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        try:
            # Remove extra whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove special characters
            text = re.sub(r'[^\w\s.,;:!?()-]', '', text)
            
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error cleaning text: {e}")
            return text
    
    async def _extract_text_from_pdf(self, file_path: Union[str, Path]) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            
            # First pass with pdfplumber for better text and table extraction
            async with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                    
                    # Extract tables if enabled
                    if self.extract_tables:
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                table_text = "\n".join(
                                    " | ".join(str(cell) for cell in row)
                                    for row in table
                                )
                                text += f"\nTable:\n{table_text}\n"
            
            # Second pass with PyPDF2 for additional text
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read()
                pdf_reader = PyPDF2.PdfReader(content)
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text not in text:  # Avoid duplicates
                        text += page_text + "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF: {e}")
            return ""
    
    async def _extract_text_from_docx(self, file_path: Union[str, Path]) -> str:
        """Extract text from DOCX file."""
        try:
            text = ""
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract tables if enabled
            if self.extract_tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        text += row_text + "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    async def _extract_text_from_txt(self, file_path: Union[str, Path]) -> str:
        """Extract text from TXT file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return (await f.read()).strip()
        except Exception as e:
            self.logger.error(f"Error extracting text from TXT: {e}")
            return ""
    
    async def process_document(
        self,
        file_path: Union[str, Path],
        custom_processor: Optional[Callable] = None
    ) -> List[Document]:
        """
        Process a document and extract text.
        
        Args:
            file_path: Path to the document
            custom_processor: Optional custom processor function
            
        Returns:
            Extracted text
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Use custom processor if provided
            if custom_processor:
                return custom_processor(file_path)
            
            # Process based on file type
            if file_path.suffix.lower() == ".pdf":
                text = await self._extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() == ".docx":
                text = await self._extract_text_from_docx(file_path)
            elif file_path.suffix.lower() == ".txt":
                text = await self._extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create documents with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "chunk": i,
                        "total_chunks": len(chunks)
                    }
                ))
            
            return documents
                
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {e}")
            return []
    
    async def process_and_store_document(
        self,
        file_path: Union[str, Path],
        vector_store: Any,
        metadata: Optional[Dict] = None,
        custom_processor: Optional[Callable] = None
    ) -> bool:
        """
        Process a document and store it in the vector store.
        
        Args:
            file_path: Path to the document
            vector_store: Vector store instance
            metadata: Optional document metadata
            custom_processor: Optional custom processor function
            
        Returns:
            True if successful, False otherwise
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Generate file hash
            file_hash = self._get_file_hash(file_path)
            
            # Prepare metadata
            doc_metadata = {
                "file_path": str(file_path),
                "file_type": file_path.suffix.lower(),
                "file_hash": file_hash,
                "processed_at": time.time(),
                **(metadata or {})
            }
            
            # Process document
            documents = await self.process_document(file_path, custom_processor)
            if not documents:
                return False
            
            # Store in vector store
            return vector_store.add_document(
                doc_id=file_hash,
                text=documents,
                metadata=doc_metadata
            )
            
        except Exception as e:
            self.logger.error(f"Error processing and storing document {file_path}: {e}")
            return False
    
    def process_directory(
        self,
        directory: Union[str, Path],
        vector_store: Optional[Any] = None,
        file_filter: Optional[Callable] = None,
        custom_metadata: Optional[Dict] = None,
        custom_processor: Optional[Callable] = None
    ) -> List[str]:
        """
        Process all documents in a directory.
        
        Args:
            directory: Directory path
            vector_store: Optional vector store instance
            file_filter: Optional function to filter files
            custom_metadata: Optional metadata to add to documents
            custom_processor: Optional custom processor function
            
        Returns:
            List of processed file paths
        """
        try:
            directory = Path(directory)
            processed_files = []
            
            for file_path in directory.glob("*"):
                if file_path.is_file():
                    # Apply file filter if provided
                    if file_filter and not file_filter(file_path):
                        continue
                    
                    # Process document
                    if vector_store:
                        success = self.process_and_store_document(
                            file_path,
                            vector_store,
                            custom_metadata,
                            custom_processor
                        )
                    else:
                        text = self.process_document(file_path, custom_processor)
                        success = bool(text)
                    
                    if success:
                        processed_files.append(str(file_path))
            
            return processed_files
            
        except Exception as e:
            self.logger.error(f"Error processing directory {directory}: {e}")
            return []
    
    def process_selected_files(
        self,
        file_paths: List[Union[str, Path]],
        vector_store: Optional[Any] = None,
        custom_metadata: Optional[Dict] = None,
        custom_processor: Optional[Callable] = None
    ) -> List[str]:
        """
        Process selected files.
        
        Args:
            file_paths: List of file paths
            vector_store: Optional vector store instance
            custom_metadata: Optional metadata to add to documents
            custom_processor: Optional custom processor function
            
        Returns:
            List of processed file paths
        """
        try:
            processed_files = []
            
            for file_path in file_paths:
                file_path = Path(file_path)
                
                if not file_path.exists():
                    self.logger.warning(f"File not found: {file_path}")
                    continue
                
                # Process document
                if vector_store:
                    success = self.process_and_store_document(
                        file_path,
                        vector_store,
                        custom_metadata,
                        custom_processor
                    )
                else:
                    text = self.process_document(file_path, custom_processor)
                    success = bool(text)
                
                if success:
                    processed_files.append(str(file_path))
            
            return processed_files
            
        except Exception as e:
            self.logger.error(f"Error processing selected files: {e}")
            return []
    
    def find_similar_documents(
        self,
        query: Union[str, List[str]],
        threshold: float = 0.8,
        doc_ids: Optional[List[str]] = None,
        filter_metadata: Optional[Dict] = None,
        **kwargs
    ) -> List[Dict]:
        """Find documents similar to the query with flexible filtering."""
        try:
            return self.vector_store.similarity_search(
                query=query,
                threshold=threshold,
                doc_ids=doc_ids,
                filter_metadata=filter_metadata,
                **kwargs
            )
        except Exception as e:
            self.logger.error(f"Error finding similar documents: {e}")
            return []
    
    def get_document_chunks(
        self,
        doc_id: str,
        chunk_indices: Optional[List[int]] = None
    ) -> List[Dict]:
        """Get specific chunks for a document."""
        try:
            return self.vector_store.get_document_chunks(doc_id, chunk_indices)
        except Exception as e:
            self.logger.error(f"Error getting document chunks for {doc_id}: {e}")
            return []
    
    def update_document_metadata(self, doc_id: str, metadata: Dict) -> bool:
        """Update metadata for a specific document."""
        try:
            return self.vector_store.update_document_metadata(doc_id, metadata)
        except Exception as e:
            self.logger.error(f"Error updating metadata for {doc_id}: {e}")
            return False
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document from the vector store."""
        try:
            return self.vector_store.delete_document(doc_id)
        except Exception as e:
            self.logger.error(f"Error deleting document {doc_id}: {e}")
            return False
    
    def get_stats(self) -> Dict:
        """Get statistics about the document processor."""
        try:
            return self.vector_store.get_stats()
        except Exception as e:
            self.logger.error(f"Error getting stats: {e}")
            return {}
    
    def cleanup(self, max_age: Optional[int] = None):
        """Clean up old documents and cache."""
        try:
            self.vector_store.cleanup(max_age)
        except Exception as e:
            self.logger.error(f"Error during cleanup: {e}") 